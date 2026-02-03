"""
Timesheet Document Generation Service
Fills in the official Florida VR/DOE timesheet template
"""
from io import BytesIO
from datetime import date, time
from typing import Optional, List
import os
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Path to the template file (use original with content controls)
TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    'templates',
    'timesheet_template.docx'
)


class TimesheetDocGenerator:
    """Generates filled timesheet documents from the official template"""

    def _format_time(self, t) -> str:
        """Format time object to string"""
        if t is None:
            return ""
        if isinstance(t, str):
            return t
        return t.strftime("%I:%M %p")

    def _format_date(self, d) -> str:
        """Format date object to string"""
        if d is None:
            return ""
        if isinstance(d, str):
            return d
        return d.strftime("%m/%d/%Y")

    def _format_date_short(self, d) -> str:
        """Format date for table (mm/dd)"""
        if d is None:
            return ""
        if isinstance(d, str):
            from datetime import datetime
            d = datetime.strptime(d, '%Y-%m-%d').date()
        return d.strftime("%m/%d/%y")

    def _fill_all_content_controls(self, doc, info_values: dict, entries: List[dict],
                                    total_hours: float, participant_name: str,
                                    signature_date):
        """
        Fill all content controls in the document.
        This method iterates through all SDT elements and fills them based on their position/context.
        """
        # Track which content control we're at for each category
        info_control_idx = 0
        time_entry_control_idx = 0

        # Expected order of info controls (based on template structure)
        info_fields = [
            info_values.get('participant_name', ''),
            info_values.get('case_id', ''),
            info_values.get('employer_name', ''),
            info_values.get('worksite_name', ''),
            info_values.get('job_title', ''),
            info_values.get('supervisor_name', ''),
            info_values.get('employer_address', ''),
            info_values.get('worksite_phone', ''),
        ]

        # Build time entry values - 6 fields per entry (date, time in, lunch out, lunch in, time out, hours)
        time_fields = []
        for entry in entries:
            entry_date = entry.get('date')
            if isinstance(entry_date, str):
                from datetime import datetime
                entry_date = datetime.strptime(entry_date, '%Y-%m-%d').date()

            time_fields.extend([
                self._format_date_short(entry_date),
                self._format_time(entry.get('start_time')),
                self._format_time(entry.get('lunch_out')),
                self._format_time(entry.get('lunch_in')),
                self._format_time(entry.get('end_time')),
                f"{entry.get('hours', 0):.1f}" if entry.get('hours', 0) > 0 else '',
            ])

        # Signature fields - signature, date, printed name
        signature_fields = [
            f"{participant_name} (signed digitally)",  # Signature
            self._format_date(signature_date),         # Date
            participant_name,                           # Printed name
        ]
        signature_control_idx = 0

        # Iterate through all content controls in document order
        all_sdts = list(doc._element.iter(qn('w:sdt')))

        for sdt in all_sdts:
            sdt_content = sdt.find(qn('w:sdtContent'))
            if sdt_content is None:
                continue

            # Get current text
            current_text = ''
            text_elements = list(sdt_content.iter(qn('w:t')))
            for t in text_elements:
                if t.text:
                    current_text += t.text

            # Check if this is a placeholder that needs filling
            if 'Click or tap' in current_text or current_text.strip() == '':
                # Determine the value based on position in document
                # We need to figure out which section this control belongs to

                # Get parent to determine context
                parent = sdt.getparent()

                # Check if this is in a table
                in_table = False
                table_idx = -1
                cell_parent = parent
                while cell_parent is not None:
                    if cell_parent.tag == qn('w:tbl'):
                        in_table = True
                        # Find which table this is
                        tables_before = len(list(doc._element.iter(qn('w:tbl'))))
                        for idx, tbl in enumerate(doc._element.iter(qn('w:tbl'))):
                            if tbl == cell_parent:
                                table_idx = idx
                                break
                        break
                    cell_parent = cell_parent.getparent()

                value = ''

                if in_table:
                    if table_idx == 0:  # Info table
                        if info_control_idx < len(info_fields):
                            value = info_fields[info_control_idx]
                            info_control_idx += 1
                    elif table_idx == 1:  # Time entries table
                        if time_entry_control_idx < len(time_fields):
                            value = time_fields[time_entry_control_idx]
                            time_entry_control_idx += 1
                else:
                    # Not in a table - likely signature section
                    if signature_control_idx < len(signature_fields):
                        value = signature_fields[signature_control_idx]
                        signature_control_idx += 1

                # Set the value in the first text element, clear the rest
                if text_elements:
                    text_elements[0].text = str(value) if value else ''
                    for t in text_elements[1:]:
                        t.text = ''

    def _clear_remaining_placeholders(self, doc):
        """Remove any remaining 'Click or tap here to enter text.' placeholders"""
        for sdt in doc._element.iter(qn('w:sdt')):
            sdt_content = sdt.find(qn('w:sdtContent'))
            if sdt_content is None:
                continue

            for t in sdt_content.iter(qn('w:t')):
                if t.text and 'Click or tap' in t.text:
                    t.text = ''

    def _fill_total_hours(self, doc, total_hours: float):
        """Fill the total hours field"""
        # Look for the total hours cell in the time table
        if len(doc.tables) < 2:
            return

        table = doc.tables[1]  # Time entries table

        # Total is typically in the last row
        if len(table.rows) > 0:
            last_row = table.rows[-1]
            # Find content control in last cell
            for cell in last_row.cells:
                for sdt in cell._element.iter(qn('w:sdt')):
                    sdt_content = sdt.find(qn('w:sdtContent'))
                    if sdt_content is not None:
                        for t in sdt_content.iter(qn('w:t')):
                            if 'Click' in (t.text or '') or t.text == '':
                                t.text = f"{total_hours:.1f}"
                                return

    def generate_timesheet(
        self,
        # Participant info
        participant_name: str,
        case_id: Optional[str],
        job_title: Optional[str],
        # Worksite info
        worksite_name: Optional[str],
        supervisor_name: Optional[str],
        worksite_phone: Optional[str],
        # Timesheet data
        entries: List[dict],
        total_hours: float,
        # Signature info
        signature_base64: Optional[str] = None,
        signature_date: Optional[date] = None,
    ) -> bytes:
        """
        Generate a filled timesheet document.
        """
        # Load template
        doc = Document(TEMPLATE_PATH)

        # Fixed employer info
        employer_name = "Career Focus Inc."
        employer_address = "6013 Wesley Grove Boulevard, Suite 202, Wesley Chapel, FL 33544"

        # Prepare info values
        info_values = {
            'participant_name': participant_name,
            'case_id': case_id or '',
            'employer_name': employer_name,
            'worksite_name': worksite_name or '',
            'job_title': job_title or '',
            'supervisor_name': supervisor_name or '',
            'employer_address': employer_address,
            'worksite_phone': worksite_phone or '',
        }

        # Fill all content controls
        self._fill_all_content_controls(
            doc,
            info_values=info_values,
            entries=entries,
            total_hours=total_hours,
            participant_name=participant_name,
            signature_date=signature_date or date.today()
        )

        # Fill total hours
        self._fill_total_hours(doc, total_hours)

        # Clear any remaining placeholders
        self._clear_remaining_placeholders(doc)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


# Singleton instance
doc_generator = TimesheetDocGenerator()
